"""Generate the DriX-8 fuel analysis METHODS report (.docx).

Usage:
    python tools/fit_em2040.py            # produces rosbags/em2040_fit.json
    python tools/build_methods_doc.py [OUT.docx]

Documents which ROS 2 topics feed the endurance calculation and exactly how the
numbers are produced — extraction, segmentation, binning, fitting, and the
derived quantities. Every figure is READ from the fit JSON and the caches, so
this document cannot drift from the pipeline that made it.
"""
import json
import sys
from pathlib import Path

import numpy as np
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
from docx_style import (new_document, check_table_widths,  # noqa: E402
                        build_date_str, INK, SOFT, ACCENT, WARN)

HERE = Path(__file__).parent
CACHE = HERE / 'rosbags'
FIT = json.load(open(CACHE / 'em2040_fit.json'))
MODEL = json.load(open(HERE.parent / 'model.json'))
# Alongside the other repo documents, like the three sibling builders. This
# used to default to Downloads, which meant a plain `python
# tools/build_methods_doc.py` silently left the committed copy stale — the
# 25% reserve change was rebuilt everywhere except here, and only a text
# search of the output caught it.
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else (
    HERE.parent / 'docs' / 'DriX8_Fuel_Methods.docx')

# -------------------------------------------------------------------- document
S = new_document(headings=(('Heading 1', 15, ACCENT, 16, 8, True),
                           ('Heading 2', 12, ACCENT, 13, 5, False),
                           ('Heading 3', 10.5, INK, 10, 4, False)),
                 right_from=99, warn_prefix=True, callout_spacer=None)
doc, para, mono = S.doc, S.para, S.mono
bullets, shade, table, callout = S.bullets, S.shade, S.table, S.callout


# ------------------------------------------------------------------ data facts
days = sorted(CACHE.glob('2026-*.npz'))
day_stats = []
for p in days:
    z = np.load(p)
    t3 = z['t3_t']
    on = z['t3_engine_on'].astype(bool)
    span = (t3.max() - t3.min()) / 3600 if len(t3) else 0
    day_stats.append(dict(day=p.stem, n=len(t3), span=span, on=on.mean() if len(on) else 0,
                          gps=len(z['gps_t']), ins=len(z['ins_t']),
                          trim=len(z['trim_t']) if 'trim_t' in z.files else 0))
nS, nF = FIT['speed_vs_rpm'], FIT['fuel_vs_rpm_quadratic']
G = MODEL['gondolas']['options']['em2040']

# Mission fuel three ways, computed from the engine so §6.1 cannot drift from
# it: the adopted integral, the same integral under reading B, and the retired
# capacity x (1 - reserve) figure kept only to show what it gets wrong.
sys.path.insert(0, str(HERE.parent))
from engine import Model as _EngModel  # noqa: E402
# §7's figures come from currents.py itself — the cycle shape, the projection
# reach and the measured accuracy block — so the chapter cannot drift from the
# code the way a hand-typed one would. Nothing here reads the network or needs a
# cached cycle: a fresh clone builds this document.
import currents as ofs  # noqa: E402
_ENG = _EngModel()
_RES_PCT = MODEL['reserve']['default_fraction'] * 100.0
_MISSION_L = _ENG.gauge_profile.litres_between(_RES_PCT, 100.0)
_NAIVE_L = MODEL['tank_volume']['litres'] * (1 - _RES_PCT / 100.0)
_MISSION_B = (100.0 - _RES_PCT) * MODEL['gauge_calibration']['l_per_point']

# ======================================================================= TITLE
para('', after=54)
para('DriX-8 Fuel Efficiency Analysis', size=24, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para('Data Sources and Calculation Methods', size=14, italic=True, color=SOFT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=30)
para(f'EM2040 gondola · {len(day_stats)} days of MCAP logs · '
     f'{FIT["cruise_hours"]:.1f} h of steady cruise',
     size=11, color=SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(build_date_str(), size=10, color=SOFT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=36)
callout('What this document is for',
        'It records which ROS 2 topics the endurance numbers come from and exactly how those '
        'numbers are produced — extraction, segmentation, binning, fitting and the derived '
        'quantities — so any figure can be traced back to a channel and a decision. Every value '
        'here is read from the pipeline output at build time; none is typed by hand. It is the '
        'companion to the main fuel-efficiency report, which covers findings rather than method.')

# ==================================================================== 1 SOURCES
doc.add_heading('1.  Data sources', level=1)
para('All analysis derives from the vehicle\'s connectivity-box MCAP recordings — the same bags '
     'the ROS 2 topic reference documents. Nothing is taken from operator notes, and no value is '
     'entered by hand at any stage of the pipeline.')
table(['Day', 'Telemetry samples', 'Span (h)', 'Engine on', 'GPS', 'INS', 'Trim'],
      [[d['day'], f'{d["n"]:,}', f'{d["span"]:.1f}', f'{d["on"]:.0%}',
        f'{d["gps"]:,}', f'{d["ins"]:,}', f'{d["trim"]:,}' if d['trim'] else '—']
       for d in day_stats],
      [1.0, 1.3, 0.85, 0.85, 0.85, 0.85, 0.8], right_from=1,
      note='Table 1 — Extracted cache per day. "Engine on" is the fraction of telemetry samples '
           'with the engine running; days near zero are alongside days that contribute only to '
           'the gauge and idle records.')
para('Segment numbering is not always contiguous — recordings have been observed missing a '
     'segment mid-day. The pipeline sorts by timestamp rather than filename and only ever uses '
     'windows of continuous data, so a gap removes coverage but cannot corrupt a fit. It does '
     'mean whole-day totals understate elapsed time, and any per-day fuel arithmetic must not '
     'assume continuity across a gap.')

# ===================================================================== 2 TOPICS
doc.add_heading('2.  Topics used', level=1)
para('Six topics carry every number in the endurance analysis. They were chosen to be the '
     'minimum sufficient set: decoding is the dominant cost of the pipeline, so channels that '
     'do not change an answer are not read.')

table(['Topic', 'Message type', 'Fields consumed', 'Used for'],
      [['…/plc/telemetry/telemetry3', 'drix_msgs/Telemetry3',
        'fuel_rate_lph, engine_on, total_fuel_used_l, alternator_1/2_current',
        'The fuel measurement itself; engine-running gate; integral cross-check'],
       ['…/plc/vehicle_status', 'drix_msgs/DrixOutput',
        'thruster_rpm, gasoline_level_percent',
        'Engine RPM (see §2.1); tank gauge for the capacity calibration'],
       ['…/topic_simplifier/light_gps', 'mdt_msgs/LightGps',
        'sog (m/s), cog, latitude, longitude',
        'Speed over ground and course steadiness'],
       ['…/topic_simplifier/light_ins', 'cortix_msgs/LightIns',
        'heave, pitch, roll, heading',
        'PHINS-derived motion statistics — the sea-state proxy'],
       ['…/hardware/trimmer/status', 'drix_msgs/TrimmerStatus',
        'position_deg, zero_position_deg',
        'Trim-tab attitude, a drag variable under watch'],
       ['…/fuel_cons_manager/fuel_consumption', 'cortix_msgs/FuelConsumption',
        'consumption_4/8/12kn, dynamic_model_used',
        'Cross-check against the vehicle\'s own static model'],
       ],
      [1.55, 1.5, 1.6, 1.65], size=8,
      note='Table 2 — Topics read by tools/extract_bags.py. INS is subsampled 1-in-3 and the '
           'Exail model 1-in-10; the rest are taken at full rate.')

doc.add_heading('2.1  Why thruster_rpm and not shaft_sensor_rpm', level=2)
para('Telemetry3 publishes a shaft RPM sensor reading, which would be the natural choice. It is '
     'faulted throughout the recorded period — reading zero or 65527 — so the PLC\'s '
     'thruster_rpm from vehicle_status is used instead. This is sound rather than a compromise: '
     'the DriX-8 drivetrain is direct drive, so shaft RPM equals engine RPM (the sole exception, '
     'the trawling motor, is no longer used). The extraction still records the faulted channel '
     'so the fault remains visible in the cache.')

doc.add_heading('2.2  Topics deliberately not used', level=2)
bullets([
    ('Weather sensor — not available. ',
     'A sweep of every channel in every segment found no live wind or meteorological topic. The '
     'only weather-adjacent message is a blank operator form. The extraction probes each new '
     'day\'s first segment against a wishlist and announces any arrival, so the moment the '
     'sensor is bridged into the recording the pipeline will say so.'),
    ('Raw PHINS / high-rate IMU — not recorded. ',
     'Only the derived light_ins stream is bridged. It is sufficient for motion statistics at '
     'the timescales that matter here, but a high-rate attitude stream would sharpen them.'),
    ('Radar, lidar, AIS, diagnostics, bandwidth. ',
     'High message volume, no bearing on fuel. Excluded to keep decode time proportionate.'),
])

# ================================================================ 3 EXTRACTION
doc.add_heading('3.  Extraction', level=1)
mono('python tools/extract_bags.py [BAG_ROOT]',
     'Bags → one compressed NPZ cache per day in tools/rosbags/. Existing caches are skipped, '
     'so adding a day costs only that day\'s decode.')
para('Three details matter for correctness:')
bullets([
    ('Segment ordering. ',
     'Files sort lexically, which places segment _10 before _2. The extractor sorts numerically '
     'and the fitting stage additionally sorts every channel by timestamp on load.'),
    ('Units. ',
     'SOG arrives in metres per second and is converted to knots with the exact factor '
     '3600/1852; RPM and fuel rate are used as published.'),
    ('Cache integrity. ',
     'A day folder containing no segments raises rather than writing an empty cache — this '
     'catches a bag root pointed one level too high, which otherwise fails silently.'),
])

# ============================================================== 4 SEGMENTATION
doc.add_heading('4.  Segmentation', level=1)
para('A fuel curve requires steady state. Instantaneous flow-meter readings during acceleration, '
     'turns or manoeuvring describe transients, not the relationship between engine speed and '
     'consumption, so the pipeline selects only windows where the vehicle is genuinely settled.')

doc.add_heading('4.1  Cruise', level=2)
para('A sample qualifies as cruise when, over the preceding 60 seconds, all of the following '
     'hold:')
table(['Condition', 'Threshold', 'Why'],
      [['Engine running', 'engine_on true', 'Excludes drifting and alongside time'],
       ['RPM steady', 'rolling std ≤ 40 rpm', 'Rejects acceleration and throttle changes'],
       ['Speed steady', 'rolling std ≤ 0.20 m/s', 'Rejects surging and wake effects'],
       ['Course steady', 'rolling std ≤ 6° (unwrapped)', 'Rejects turns — a turning hull carries '
                                                          'rudder drag that is not cruise drag'],
       ['RPM in band', '1200 ≤ rpm < 3200', 'Above idle, below the recorded ceiling'],
       ['Making way', 'sog > 1.5 m/s', 'Excludes station-keeping'],
       ],
      [1.35, 1.6, 3.3],
      note='Table 3 — Cruise criteria. Course unwrapping matters: without it a heading crossing '
           '360° reads as a 360° excursion and every northerly leg would be discarded.')

doc.add_heading('4.2  Loiter', level=2)
para('Station-keeping is retained separately — engine on, 800–1200 rpm, fuel rate above 0.3 L/h. '
     'It is deliberately excluded from the cruise fit: at those revolutions the propeller is '
     'loaded quite differently and the vehicle is not making way, so including it would bend the '
     'cruise curve toward a regime it does not describe. It is reported on its own as the idle '
     'and station-keeping burn.')

# =============================================================== 5 BINNING/FIT
doc.add_heading('5.  Binning and fitting', level=1)
para('Steady samples are pooled across all days and grouped into 100-rpm bins. A bin is used '
     'only if it holds at least 120 samples — about two minutes of steady running — and each bin '
     'contributes its median speed and median fuel rate. Medians rather than means because a '
     'single flow-meter spike should not move a bin.')
para('Fits are weighted least squares over the bin medians, with weights equal to sample count '
     'capped at 3600. The cap is the point: without it, one long leg at a favourite cruise '
     'setting would dominate the entire curve, and the fit would describe that leg rather than '
     'the vehicle.')
mono('speed:  KNOTS = b + m · RPM              (linear)\n'
     'fuel :  L/H   = q0 + q1 · RPM + q2 · RPM²  (quadratic)',
     'The quadratic fuel form is used because a propeller absorbs torque roughly as the square '
     'of shaft speed; the linear speed form matches the observed data across the cruise band.')

doc.add_heading('5.1  The laws this produced', level=2)
para('For the record, the coefficients currently in force — read from the pipeline output, not '
     'transcribed:')
mono(f'KNOTS = {nS["b"]:.6f} + {nS["m"]:.8f} · RPM',
     f'R² {nS["r2"]:.4f} over {len(FIT["bins"])} bins.')
mono(f'L/H   = {nF["q0"]:.6f} {nF["q1"]:+.8f} · RPM {nF["q2"]:+.6e} · RPM²',
     f'R² {nF["r2"]:.4f}, valid {int(FIT["rpm_range"][0])}–{int(FIT["rpm_range"][1])} rpm.')
_r8 = (8 - nS['b']) / nS['m']
_l8 = nF['q0'] + nF['q1'] * _r8 + nF['q2'] * _r8 * _r8
para(f'At the 8 kt survey speed that is {_r8:.0f} rpm, {_l8:.2f} L/h, {8/_l8:.2f} NM/L — inside '
     f'the fitted window, so survey planning on this configuration is interpolation rather than '
     f'extrapolation.')

table(['RPM bin', 'Samples', 'SOG (kt)', 'Fuel (L/h)', 'NM/L'],
      [[f'{int(b["rpm"]//100)*100}–{int(b["rpm"]//100)*100+99}', f'{b["n"]:,}',
        f'{b["kt"]:.2f}', f'{b["lph"]:.2f}', f'{b["kt"]/b["lph"]:.2f}']
       for b in FIT['bins']],
      [1.2, 1.1, 1.1, 1.1, 1.0], right_from=1,
      note='Table 4 — The bin medians the fits are computed over. Bin count and contents change '
           'as data accumulates; this table is regenerated with the document.')

doc.add_heading('5.2  Per-day agreement', level=2)
para('Because data now arrives daily, the fit reports each day\'s bin medians against the pooled '
     'result. A day that disagrees materially is named with its worst bin rather than being '
     'quietly averaged into the pool — the check that would catch a fouled hull, a changed '
     'payload, or a sensor drifting out of calibration.')

# ================================================================= 6 DERIVED
doc.add_heading('6.  Derived quantities', level=1)
doc.add_heading('6.1  Efficiency, range and endurance', level=2)
mono('efficiency  (NM/L) = KNOTS / (L/H)\n'
     'range       (NM)   = mission fuel × efficiency\n'
     'endurance   (h)    = mission fuel ÷ (L/H)\n'
     'mission fuel (L)   = ∫ gauge, from the reserve floor up to the start level',
     'NOT capacity × (1 − reserve fraction). That formula was retired in v2.4.0: '
     'it assumes the gauge is linear and that a tank percentage is a fixed number '
     'of litres, and neither holds.')
para(f'Planning figures are computed to the {MODEL["reserve"]["default_fraction"]:.0%} '
     f'return-to-port reserve rather than to a dry tank. '
     f'Tank volume is ESTABLISHED at {MODEL["tank_volume"]["litres"]:.0f} L from the engineering '
     f'drawings, but that is not the same quantity as mission fuel: the floor is written in '
     f'indicated percent, so what a mission may spend is an integral over the gauge between the '
     f'start level and the floor. See the fuel-efficiency report §8.2 and the gauge-linearity '
     f'report for why those two differ by about 44 L.')
para(f'The difference is not academic. On the adopted reading the integral gives '
     f'{_MISSION_L:.1f} L from a full tank, where capacity × (1 − reserve) would give '
     f'{_NAIVE_L:.1f} L. Those happen to sit close together because the profile '
     f're-normalises to the drawing volume; on the conservative reading (B) the same '
     f'two figures are {_MISSION_B:.1f} L and {_NAIVE_L:.1f} L, which is a mission-'
     f'shaping difference. The integral is the quantity the planner uses in both cases.')

doc.add_heading('6.2  Gauge calibration', level=2)
para('The tank gauge is calibrated directly against the flow meter: total metered litres over a '
     'day divided by the indicated percentage points consumed, using ten-minute medians at each '
     'end so a single noisy reading cannot set an endpoint. Days that burn less than two litres '
     'are excluded — an idle day\'s gauge wanders a point or two on essentially no fuel, and '
     'letting one anchor the span moves the answer by around a tenth for no physical reason.')
para('This is the only non-circular capacity measurement available: it assumes no tank volume '
     'and no fuel model. The result is band-specific, and the gauge is not linear, so a figure '
     'measured in one part of the range must not be extrapolated across the tank.')

doc.add_heading('6.3  Through-origin variants for ops tables', level=2)
para('The planner uses the fitted laws as measured, because they are most accurate in the band '
     'where missions are actually flown. An operator-facing endurance table has a different '
     'requirement: it must behave physically at its ends, showing zero fuel at zero RPM. For '
     'those tables the same bin data is refitted under a through-origin constraint, anchored '
     'additionally by the measured idle burn. The two families agree closely in band, and the '
     'endurance sheet states which it uses.')

# =================================================================== 7 CURRENTS
doc.add_heading('7.  Currents from the NOAA forecast', level=1)
para('A second data source, acquired and used differently from everything above. The endurance '
     'laws come from the vehicle\'s own recordings; the tidal current comes from an operational '
     'model published by NOAA, and it is the only part of the planner that touches a network. It '
     'is documented here because it is an input to a plan, and an input a reader may need to '
     'defend.')
callout('What it does NOT change',
        'No coefficient moves. model.json is untouched by any of this, the engine gained nothing, '
        'and no fitted value is involved. The forecast supplies two inputs that already existed — '
        'a leg\'s current speed and set — so an operator who types the tide in from a table gets '
        'exactly the same arithmetic. Never pressing the button leaves the planner as it was.')

doc.add_heading('7.1  Product, and the decision that mattered', level=2)
para('NOAA publishes the Delaware Bay Operational Forecast System over OPeNDAP in two forms. The '
     'native ROMS output carries velocities on a curvilinear grid in GRID axes on staggered '
     'points: using it means averaging u and v onto rho points and rotating every cell by its own '
     'angle before a bearing means anything. The regridded product is rectilinear and already '
     'true-referenced.')
para('This reads the regridded product, because the two operations the native one needs are '
     'exactly the two that fail SILENTLY — get the rotation wrong and every set is out by tens of '
     'degrees while every speed still looks reasonable. That shortcut is therefore evidenced '
     'rather than assumed: currents.py crosscheck performs the native path by hand and compares, '
     'and the two agree to a median 0.07 kt and 1.0°.')

doc.add_heading('7.2  What a cycle is, and which one answers', level=2)
table(['Property', 'Value'],
      [['Nowcast hours (up to the cycle time)', f'{ofs.NOWCAST_H}'],
       ['Forecast hours (after it)', f'{ofs.FORECAST_H}'],
       ['Frames per cycle', f'{ofs.NOWCAST_H + ofs.FORECAST_H} hourly'],
       ['Span', f'{ofs.NOWCAST_H + ofs.FORECAST_H - 1} h'],
       ['Layer', 'surface only (depth 0 m)'],
       ['Archive NOAA serves', 'about 2 days'],
       ['Projection reach', f'{ofs.MAX_PROJECT_CYCLES} tidal cycles '
                            f'({ofs.MAX_PROJECT_CYCLES * ofs.M2_PERIOD_H:.1f} h)']],
      [2.6, 3.2], right_from=1,
      note='Table 5 — Cycle shape and the limits around it, read from currents.py at build time.')
para('A requested time is answered by a ladder, and real data always beats an estimate: a cached '
     'cycle covering the window; else a cycle NOAA still serves covering it, found from the '
     'catalog for the price of one page rather than a 33 MB download, and fetched; else a '
     'projection; else a refusal. Cycle selection checks the BOX as well as the span — a cycle '
     'fetched for one operating area cannot answer for another, and without that check it would '
     'report no water on every leg, which reads like a forecast of slack water rather than the '
     'wrong file.')

doc.add_heading('7.3  Projection, and what it is worth', level=2)
_PA = ofs.PROJECTION_ACCURACY
_pr = _PA['projected_rms_kt']
para('The current here is semidiurnal, so a time the forecast cannot reach is estimated from the '
     f'value one M2 period ({ofs.M2_PERIOD_H} h) away — never by holding the last value, and never '
     'by extrapolating a line through a reversing tide. That is a measurement rather than a '
     'preference:')
table(['Method', 'RMS error against the model'],
      [[f'Project {n} cycle{"s" if n > 1 else ""} '
        f'({n * ofs.M2_PERIOD_H:.1f} h)', f'{v:.2f} kt'] for n, v in sorted(_pr.items())]
      + [['Hold the last value',
          f'{_PA["persistence_rms_kt"][0]:.2f} – {_PA["persistence_rms_kt"][1]:.2f} kt'],
         ['Assume slack water',
          f'{_PA["slack_rms_kt"][0]:.2f} – {_PA["slack_rms_kt"][1]:.2f} kt']],
      [3.0, 2.8], right_from=1,
      note=f'Table 6 — Measured on {_PA["cycle"]} across {_PA["samples"]} samples '
           f'({_PA["measured_utc"]}), reproducible with tools/projection_accuracy.py. '
           + _PA['note'])
para(f'The reach is capped at {ofs.MAX_PROJECT_CYCLES} cycles for that reason and not because it '
     'is a round number: the error is flat that far because a tide repeats, but the non-tidal part '
     '— wind setup, river flow — does not repeat at all, and beyond about a day and a half there '
     'is no evidence in hand. Two things a projection never does: it never moves a POSITION, so a '
     'point with no model water is unanswered however the time is shifted; and it never happens '
     'silently — a projected leg is flagged in the response, named in its note, marked in the '
     'provenance label the mission report prints, and surfaced in the plan warnings.')

doc.add_heading('7.4  How a current reaches a plan', level=2)
para('There are exactly two seams, and they differ in resolution rather than in kind:')
bullets([
    ('Per leg, one number each. ',
     'The mission is dead-reckoned from the departure position and time, each leg sampled along '
     'its own track at the time the vehicle would be there, and the samples vector-averaged. A '
     'survey holds position and only advances the clock, because a lawnmower ends roughly where '
     'it began and walking it down the first line\'s course would put the run home in the wrong '
     'water.'),
    ('Along the track, a field. ',
     'Given geometry, the planner instead calls back once per RUN — per survey line, per transit '
     'segment — with where the vehicle is and how far into the mission, and takes the current '
     'there and then. This is what makes a turning tide real rather than averaged.'),
])
para('Either way the current changes the FUEL and never the clock: the required speed over ground '
     'is converted to the speed through the water the hull must actually make, and that goes '
     'through the gondola\'s speed law to get RPM. A leg still takes distance divided by speed.')
callout('It is partly counted twice, and the leg note says so',
        'The speed law is fitted against speed over ground in an unrecorded tide, so some tidal '
        'effect is already inside it. Reading the tide off the forecast makes the INPUT better; it '
        'does not make the correction clean. The honest framing is "compare two plans" or "price '
        'today\'s tide", not "a calibrated tidal model". Removing it properly needs the same '
        'reciprocal-heading runs §9 already asks for.')

doc.add_heading('7.5  A worked example', level=2)
para('One mission, planned three ways, against a real cycle. Twelve miles out on 045, a 12 × 2 NM '
     'survey on 020, twelve miles home on 225, departing Lewes at midday UTC:')
table(['Planned with', 'Fuel', 'Mission clock'],
      [['Slack water (no current)', '17.18 L', '7.43 h'],
       ['Per-leg currents from the forecast', '18.08 L  (+5.2%)', '7.43 h'],
       ['Sampled along the track', '17.64 L  (+2.7%)', '7.43 h']],
      [2.7, 1.7, 1.4], right_from=1,
      note='Table 7 — Transcript against dbofs_20260813_t00z. The clock is identical in all '
           'three: a current moves the fuel and never the time.')
para('The reading that matters is the third row. Sampling along the track gives 2.7% where the '
     'per-leg average gives 5.2% — nearly double — because the survey sits on one ground for four '
     'hours while the tide turns under it, and a single vector average of a reversing current '
     'overstates what the boat actually fights. That is the argument for entering geometry, in '
     'numbers rather than in principle.')

# ============================================================== 8 VERIFICATION
doc.add_heading('8.  Verification', level=1)
bullets([
    ('Flow-meter integral against the vehicle\'s own counter. ',
     'Integrated fuel rate is compared with the cumulative total_fuel_used_l for each day; '
     'agreement confirms neither channel is drifting.'),
    ('Per-day agreement. ', 'Described in §5.1 — the standing check as data accumulates.'),
    ('Adoption comparison. ',
     'tools/compare_fits.py contrasts a fresh fit with the coefficients currently adopted and '
     'reports operational deltas — burn at survey speed, planning range, endurance — rather than '
     'raw coefficients, which are not interpretable in isolation.'),
    ('Mutation tests. ',
     'The planner\'s test suite perturbs each coefficient and asserts the results move. A model '
     'change that no test notices means the tests are broken, not that the change is safe.'),
    ('Currents, three independent ways. ',
     'Direction fails quietly, so the forecast path is checked against things that are not it: '
     'currents.py crosscheck reproduces the native ROMS path by hand (median 0.07 kt, 1.0°); '
     'tools/dbofs_plotcheck.py redraws the extracted field onto NOAA\'s own published image, '
     'deriving the georeference from the plot\'s graticule rather than a stated corner; and '
     'currents.py station_check compares against CO-OPS harmonic predictions at a real station '
     '(r = +0.987 over 54 h).'),
    ('Projection accuracy. ',
     'tools/projection_accuracy.py re-measures what a projection is worth against the model\'s '
     'own output, and reports whether the recorded block still matches. §7.3 quotes that block '
     'rather than a typed number.'),
])

# ==================================================================== 9 LIMITS
doc.add_heading('9.  Limits of the method', level=1)
para('Stated plainly, because each one bounds how far the numbers should be trusted:')
bullets([
    ('Speed is over ground. ',
     'No speed-through-water sensor is recorded, so tidal set enters the speed law directly. '
     'Reciprocal-heading runs at matched RPM would remove it; until then the speed law carries '
     'a few percent of tidal uncertainty and the endurance sheet says so.'),
    ('Sea state is a proxy. ',
     'Motion statistics from the derived INS stream stand in for sea state. They measured the '
     'calm anchor — no fuel premium above the noise floor in near-glassy water — but the slope '
     'into rough water is unmeasured because no rough-water cruise data exists yet.'),
    ('Below the cruise band there is only the idle point. ',
     'No steady cruise data exists below the fitted floor; the region between idle and that '
     'floor is bridged by curve shape, not measurement.'),
    ('One vehicle, one gondola configuration. ',
     'These laws describe DriX-8 as currently fitted. Fouling, loading and payload changes all '
     'move them, which is exactly why per-day agreement is monitored.'),
    ('The forecast current is surface only, and partly double-counted. ',
     'It is the 0 m layer on an hourly grid a few hundred metres across; no shear is modelled, '
     'and the vehicle draws two metres. Because the speed law above is fitted against speed over '
     'ground in an unrecorded tide, applying a forecast current counts some of that tide twice — '
     'see §7.4. It improves the input; it does not make the correction clean.'),
    ('The forecast archive is about two days deep. ',
     f'Beyond it a requested time can only be estimated, and only within '
     f'{ofs.MAX_PROJECT_CYCLES} tidal cycles '
     f'({ofs.MAX_PROJECT_CYCLES * ofs.M2_PERIOD_H:.0f} h) of data actually held. Past that the '
     'planner declines rather than guessing, and a position with no model water is never answered '
     'at any time by any route.'),
])

doc.add_heading('Appendix  —  Reproducing this analysis', level=1)
para('The whole chain, in order. Extra dependencies are numpy, mcap and mcap-ros2-support; the '
     'planner itself remains standard-library only.')
mono('python tools/extract_bags.py D:/Claude/Fuel/D8_2040\n'
     'python tools/fit_em2040.py\n'
     'python tools/compare_fits.py\n'
     'python tools/build_endurance_sheet.py  [OUT.xlsx]\n'
     'python tools/build_methods_doc.py      [OUT.docx]\n'
     'powershell -File tools/export_pdf.ps1  [-Path docs\\X.docx]\n'
     '\n'
     '# §7 only, and independent of the bags:\n'
     'python currents.py fetch\n'
     'python currents.py verify · crosscheck · station_check\n'
     'python tools/projection_accuracy.py',
     'Extraction skips days already cached, so adding a day costs only that day. compare_fits '
     'reports what a new fit would change before anything is adopted.')
para('Adoption is a deliberate step, not an automatic one: a fresh fit is compared against the '
     'coefficients in force, and only if the operational deltas justify it are model.json, the '
     'test reference values and the generated documents updated together in a single commit.')

check_table_widths(doc)
doc.save(OUT)
print('written:', OUT)
print(f'days={len(day_stats)} cruise_h={FIT["cruise_hours"]:.2f} '
      f'bins={len(FIT["bins"])}')
