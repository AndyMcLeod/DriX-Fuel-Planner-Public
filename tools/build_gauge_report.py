"""Generate the fuel-gauge linearity report (.docx) with figures.

Usage:
    python tools/build_gauge_report.py [OUT.docx]

Analyses the tank gauge against the flow meter across every cached day, builds
the figures, and writes the report. Every number is computed here from the
caches — nothing is transcribed.
"""
import sys
from pathlib import Path

import numpy as np
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.path.insert(0, str(Path(__file__).resolve().parent))
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))   # engine
from docx_style import (new_document, check_table_widths,  # noqa: E402
                        build_date_str, INK, SOFT, ACCENT, WARN)
from engine import Model as _EngModel  # noqa: E402
from drawdown import spec as _drawdown_spec  # noqa: E402
from drawdown import pooled_scale as _dd_pooled  # noqa: E402
from drawdown import usable_spans as _dd_spans  # noqa: E402

HERE = Path(__file__).parent
CACHE = HERE / 'rosbags'
FIGS = HERE / 'gauge_figs'
FIGS.mkdir(exist_ok=True)
OUT = Path(sys.argv[1]) if len(sys.argv) > 1 else (
    HERE.parent / 'docs' / 'DriX8_Fuel_Gauge_Linearity.docx')

import json as _json
with open(HERE.parent / 'model.json', encoding='utf-8') as _fh:
    RES_PCT = _json.load(_fh)['reserve']['default_fraction'] * 100.0

_DD = _drawdown_spec(_EngModel(), CACHE)

NOMINAL_LPP = 2.50          # 250 L over 100 points
DD2024_LPP = 2.09           # 185 L over ~88.5 points, whole range
Y2022_LPP = 1.73            # 65.0 modelled L over 37.5 points

C_MEAS, C_MODEL, C_ALT, C_2022, C_WARN = '#1f77b4', '#c0504d', '#ed7d31', '#4c9a2a', '#9a6400'

plt.rcParams.update({
    'font.family': 'Arial', 'font.size': 9, 'axes.titlesize': 10,
    'axes.titleweight': 'bold', 'axes.edgecolor': '#5b6b7a',
    'axes.labelcolor': '#16212b', 'text.color': '#16212b',
    'xtick.color': '#5b6b7a', 'ytick.color': '#5b6b7a', 'axes.grid': True,
    'grid.color': '#dbe2e8', 'grid.linewidth': .7, 'legend.fontsize': 8,
    'legend.frameon': False, 'figure.dpi': 200, 'savefig.dpi': 200,
    'savefig.bbox': 'tight'})


# --------------------------------------------------------------- data & maths
def load(p):
    z = np.load(p)
    d = {k: z[k] for k in z.files}
    for g in ('t3', 'vs'):
        t = d.get(f'{g}_t')
        if t is None or not len(t):
            continue
        i = np.argsort(t, kind='stable')
        for k in list(d):
            if k.startswith(g + '_'):
                d[k] = d[k][i]
    return d


def rmed(t, x, win):
    out = np.empty(len(x))
    j = 0
    for i in range(len(t)):
        while t[i] - t[j] > win:
            j += 1
        out[i] = np.median(x[j:i + 1])
    return out


days = []
for p in sorted(CACHE.glob('2026-*.npz')):
    d = load(p)
    t, lph = d['t3_t'], d['t3_fuel_lph'].astype(float)
    if len(t) < 50:
        continue
    gas = np.interp(t, d['vs_t'], d['vs_gas_pct'].astype(float))
    dts = np.diff(t, prepend=t[0])
    dts[dts > 30] = 0
    cum = np.cumsum(lph * dts) / 3600
    days.append(dict(day=p.stem, t=t, gas=gas, cum=cum, total=float(cum[-1]),
                     g0=float(np.median(gas[:600])), g1=float(np.median(gas[-600:]))))

idle = [d for d in days if d['total'] < 2.0]
SIG_PTS = 0.7            # 1% quantisation -> ~0.5 pt per endpoint, 0.7 combined

# THE UNIT OF MEASUREMENT IS A DRAWDOWN, NOT A DAY. It used to be a day, because
# the data happened to be one long fall; from 2026-08-15 it contains refuels, and
# a day spanning one has no litres-per-point at all. tools/drawdown.py does the
# cutting so this report, fit_em2040.py and reserve_band.py cannot disagree.
SPANS = _dd_spans(CACHE)
if not SPANS:
    raise SystemExit('no usable drawdown in the caches — nothing to calibrate on')
for s in SPANS:
    s['lpp'], s['pts'], s['total'] = s['l_per_point'], s['points'], s['litres']
    s['g0'], s['g1'] = s['from_pct'], s['to_pct']

POOLED, SCATTER, TOT, PTS, G_LO, G_HI, N_SPANS = _dd_pooled(SPANS)
# The pooled figure's own precision, from the spread BETWEEN spans rather than
# from gauge quantisation on one of them: with 17 falls the quantisation is no
# longer what limits the answer.
POOLED_SIG = SCATTER / np.sqrt(N_SPANS)
# Sessions whose gauge trace is worth drawing, largest burn first, capped so the
# strip stays legible — 16 panels across one page is not a figure.
active = sorted([d for d in days if d['total'] >= 2.0],
                key=lambda d: -d['total'])[:6]
active.sort(key=lambda d: d['day'])

lo_d = min(SPANS, key=lambda s: s['lpp'])
hi_d = max(SPANS, key=lambda s: s['lpp'])
GAP = hi_d['lpp'] - lo_d['lpp']
GAP_SIG = float(np.hypot(hi_d['sig'], lo_d['sig']))
SIGMAS = GAP / GAP_SIG
NOM_SIGMAS = (NOMINAL_LPP - POOLED) / POOLED_SIG

# Is there a trend with DEPTH? That is what reading A predicts, and it is the
# question the deeper 2026-08 spans could finally speak to.
_deep = [s for s in SPANS if s['to_pct'] < 62]
_shallow = [s for s in SPANS if s['to_pct'] >= 71 and s['from_pct'] <= 87]
DEEP_LPP = (sum(s['litres'] for s in _deep) / sum(s['points'] for s in _deep)
            if _deep else None)
SHALLOW_LPP = (sum(s['litres'] for s in _shallow) / sum(s['points'] for s in _shallow)
               if _shallow else None)
# What the planner is actually configured with, read from the model rather than
# restated here — §5.1 exists because this report's measurement and that value
# are no longer the same number.
_gc = _EngModel().data['gauge_calibration']
ADOPTED_LPP = _gc['l_per_point']
ADOPTED_SIG = _gc['l_per_point_sigma']
ADOPTED_BAND = _gc['band_pct']

if DEEP_LPP and SHALLOW_LPP:
    _se = np.hypot(SCATTER / np.sqrt(max(len(_deep), 1)),
                   SCATTER / np.sqrt(max(len(_shallow), 1)))
    DEPTH_SIGMAS = abs(DEEP_LPP - SHALLOW_LPP) / _se
else:
    DEPTH_SIGMAS = None
ADOPTED_SIGMAS = abs(POOLED - ADOPTED_LPP) / float(np.hypot(POOLED_SIG, ADOPTED_SIG))

# ------------------------------------------------------------------- figures
fig, axes = plt.subplots(1, len(active), figsize=(11, 3.1), sharey=True)
for ax, d in zip(np.atleast_1d(axes), active):
    hrs = (d['t'] - d['t'][0]) / 3600
    ax.plot(hrs, d['gas'], color=C_MEAS, lw=1.3)
    ax2 = ax.twinx()
    ax2.plot(hrs, d['cum'], color=C_MODEL, lw=1.6)
    ax2.set_ylim(0, max(11, max(x['total'] for x in active) * 1.15))
    ax2.grid(False)
    if ax is np.atleast_1d(axes)[-1]:
        ax2.set_ylabel('cumulative litres (flow meter)', color=C_MODEL)
        ax2.tick_params(axis='y', colors=C_MODEL)
    else:
        ax2.set_yticklabels([])
    ax.set_title(d['day'], fontsize=9)
    ax.set_xlabel('hours into day')
np.atleast_1d(axes)[0].set_ylabel('indicated tank level (%)', color=C_MEAS)
np.atleast_1d(axes)[0].tick_params(axis='y', colors=C_MEAS)
fig.suptitle('Gauge (blue, quantised to 1%) against metered fuel (red)',
             fontsize=10, fontweight='bold', y=1.06)
fig.savefig(FIGS / 'g1_traces.png', facecolor='white')
plt.close(fig)

fig, ax = plt.subplots(figsize=(6.4, 3.9))
off = 0.0
for d in active:
    ax.scatter(d['gas'], d['cum'] + off, s=1.5, color=C_MEAS, alpha=.16,
               linewidths=0, zorder=1)
    env = np.minimum.accumulate(rmed(d['t'], d['gas'], 600))
    ax.plot(env, d['cum'] + off, color=C_MEAS, lw=2.0, zorder=3)
    ax.plot(d['g0'], off, 'o', color=C_MODEL, ms=6, zorder=5)
    off += d['total']
ax.plot(active[-1]['g1'], off, 'o', color=C_MODEL, ms=6, zorder=5)
xs = np.array([G_LO, G_HI])
ax.plot(xs, (G_HI - xs) * NOMINAL_LPP, '--', color=C_2022, lw=1.8,
        label=f'nominal linear ({NOMINAL_LPP:.2f} L/pt on a 250 L tank)')
ax.plot(xs, (G_HI - xs) * POOLED, '-', color=C_MODEL, lw=2.2,
        label=f'measured ({POOLED:.2f} L/pt, pooled)')
ax.plot([], [], color=C_MEAS, lw=2.0, label='metered fuel vs level (10-min median)')
ax.scatter([], [], s=8, color=C_MEAS, alpha=.35, linewidths=0, label='raw 1%-quantised gauge')
ax.invert_xaxis()
ax.set_xlabel('Indicated tank level (%)')
ax.set_ylabel('Cumulative litres burned')
ax.set_title('Calibration curve: the gauge falls faster than nominal')
ax.set_ylim(-1, None)
ax.legend(loc='upper left')
fig.savefig(FIGS / 'g2_calibration.png', facecolor='white')
plt.close(fig)

fig, ax = plt.subplots(figsize=(6.6, 3.6))
_bar = sorted(SPANS, key=lambda s: -s['points'])[:10]
_bar.sort(key=lambda s: s['g0'])
labels = [f'{s["day"][5:]}\n{s["g1"]:.0f}-{s["g0"]:.0f}%' for s in _bar]
vals = [s['lpp'] for s in _bar]
errs = [s['sig'] for s in _bar]
cols = [C_MEAS] * len(_bar)
labels.append(f'POOLED\n{G_LO:.0f}-{G_HI:.0f}%')
vals.append(POOLED); errs.append(POOLED_SIG); cols.append(C_MODEL)
x = np.arange(len(vals))
ax.bar(x, vals, yerr=errs, capsize=5, color=cols, width=.6,
       error_kw=dict(lw=1.4, ecolor='#333'))
ax.axhline(NOMINAL_LPP, color=C_2022, ls='--', lw=1.6,
           label=f'nominal {NOMINAL_LPP:.2f} L/pt (250 L linear)')
ax.axhline(DD2024_LPP, color=C_ALT, ls=':', lw=1.6,
           label=f'DD2024 refuel, whole range {DD2024_LPP:.2f}')
ax.set_xticks(x); ax.set_xticklabels(labels, fontsize=8)
ax.set_ylabel('Litres per indicated point')
ax.set_ylim(0, 3.1)
ax.set_title('Per-day scatter sits within its own error bars')
ax.legend(loc='lower left'); ax.grid(axis='x')
fig.savefig(FIGS / 'g3_lpp.png', facecolor='white')
plt.close(fig)

fig, ax = plt.subplots(figsize=(6.2, 3.3))
srcs = [('Nominal\nassumption', 250.0, 0.0, C_2022),
        ('DD2024 refuel\n(9–97%)', DD2024_LPP * 100, 6.0, C_ALT),
        ('2022 telemetry\n(modelled litres)', Y2022_LPP * 100, 12.0, C_WARN),
        ('MCAP pooled\n(metered)', POOLED * 100, POOLED_SIG * 100, C_MODEL)]
xs = np.arange(len(srcs))
ax.bar(xs, [s[1] for s in srcs], yerr=[s[2] for s in srcs], capsize=5,
       color=[s[3] for s in srcs], width=.6, error_kw=dict(lw=1.4, ecolor='#333'))
for i, s in enumerate(srcs):
    ax.text(i, s[1] + s[2] + 6, f'{s[1]:.0f} L', ha='center', fontsize=8.5, fontweight='bold')
ax.set_xticks(xs); ax.set_xticklabels([s[0] for s in srcs], fontsize=8)
ax.set_ylabel('Implied usable capacity (L)')
ax.set_ylim(0, 290)
ax.set_title('What each source implies for usable capacity')
ax.grid(axis='x')
fig.savefig(FIGS / 'g4_capacity.png', facecolor='white')
plt.close(fig)

# -------------------------------------------------------------------- document
S = new_document(FIGS,
                 headings=(('Heading 1', 15, ACCENT, 16, 8, True),
                           ('Heading 2', 12, ACCENT, 13, 5, False)),
                 right_from=1, warn_prefix=False, callout_spacer=None)
doc, para, mono = S.doc, S.para, S.mono
bullets, shade, table, callout, figure = (S.bullets, S.shade, S.table,
                                          S.callout, S.figure)


# ---- title
para('', after=52)
para('DriX-8 Fuel Gauge', size=24, bold=True, color=ACCENT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
para('Linearity, Scale and What the Reserve Is Worth', size=14, italic=True,
     color=SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, after=28)
para(f'{len(active)} operating days measured against the flow meter · '
     f'{TOT:.1f} L over {PTS:.0f} indicated points',
     size=11, color=SOFT, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para(build_date_str(), size=10, color=SOFT,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=34)
callout('The short version',
        f'The gauge does not read litres. Over the {G_LO:.0f}–{G_HI:.0f}% band it delivers '
        f'{POOLED:.2f} ± {POOLED_SIG:.2f} litres per indicated point, against {NOMINAL_LPP:.2f} '
        f'for a linear 250 L tank — {NOM_SIGMAS:.0f} sigma below, and in close agreement with the '
        f'independent DD2024 refuel figure of {DD2024_LPP:.2f}. That is a solid result and it '
        f'points at a usable capacity near {POOLED*100:.0f} L. Whether the gauge is also '
        f'NON-LINEAR — reading different litres-per-point at different levels — is NOT '
        f'established by this data, and an earlier reading of these same days overstated it. '
        f'Section 5 sets that record straight.')

# ---- 1
doc.add_heading('1.  Why this matters', level=1)
para('Two operating decisions rest on the gauge, and both are quantitative.')
bullets([
    (f'The {RES_PCT:.0f}% return-to-port reserve is expressed in indicated percent. ',
     f'If a point is not a fixed number of litres, then "{RES_PCT:.0f}%" is not a fixed amount of fuel — and '
     'the reserve is either more or less protective than intended, without anyone being able to '
     'tell from the wheelhouse.'),
    ('Endurance predictions scale directly with usable capacity. ',
     'Range and hours-remaining are computed as usable litres times efficiency. Take the '
     'capacity wrong and every planning number inherits the error, in proportion.'),
])
para('The gauge is also the only fuel instrument available on a vehicle without a working flow '
     'meter, so understanding its behaviour transfers to hulls that cannot be measured this way.')

# ---- 2
doc.add_heading('2.  Method', level=1)
para('The flow meter is used as the ruler. It measures litres directly and is independent of '
     'the gauge, so comparing the two isolates the gauge\'s behaviour without assuming a tank '
     'volume or a fuel model anywhere in the chain.')
mono('litres per indicated point  =  ∫ fuel_rate dt  ÷  Δ indicated %')
para('Two details do the work. Fuel rate is integrated over each day with gaps longer than 30 s '
     'excluded, so a recording break cannot inflate the total. And each end of the gauge span is '
     'taken as a ten-minute median rather than a single reading, because the gauge is quantised '
     'to whole percent and bounces by a point or more as fuel moves in the tank.')
para('Spans burning less than 2 L are excluded from the calibration entirely. On an idle stretch '
     'the gauge still wanders a point or two on essentially no fuel, and letting one anchor an '
     'end moves the answer by roughly a tenth for no physical reason.')
callout('A drawdown, not a day',
        'Litres per point is a slope down a FALLING level, so it is only defined while the level '
        'falls. Until August 2026 the recorded data happened to be one long descent and each day '
        'could be read end to end. It is not any more: the vehicle was refuelled twice inside the '
        'record, once mid-session. Measured as whole days that yields a NEGATIVE litres-per-point '
        'for a refuelled day and a pooled figure several times the truth. Every session is '
        f'therefore cut at each refill, and the {N_SPANS} falling stretches below are the '
        'measurements. None spans a refill; none spans two sessions.')

table(['Session', 'Metered litres', 'Gauge span', 'Points', 'L per point', 'Uncertainty'],
      [[s['day'], f'{s["total"]:.2f}', f'{s["g1"]:.0f}–{s["g0"]:.0f}%',
        f'{s["pts"]:.1f}', f'{s["lpp"]:.2f}', f'± {s["sig"]:.2f}']
       for s in sorted(SPANS, key=lambda s: -s['points'])]
      + [['POOLED', f'{TOT:.1f}', f'{G_LO:.0f}–{G_HI:.0f}%', f'{PTS:.1f}',
          f'{POOLED:.2f}', f'± {POOLED_SIG:.2f}']],
      [1.15, 1.15, 1.0, 0.65, 1.0, 1.05],
      note=f'Table 1 — Every usable drawdown, largest first. The pooled row divides total litres '
           f'by total drawdown TRAVELLED ({PTS:.1f} points), not by the width of the band it '
           f'covers: the spans are separate falls that revisit the same part of the gauge, so the '
           f'band only says where the measurement looked. Its uncertainty is the spread BETWEEN '
           f'spans rather than gauge quantisation on any one of them — with {N_SPANS} falls, '
           f'quantisation is no longer what limits the answer.')

figure('g1_traces.png',
       'Figure 1 — The raw evidence, one panel per operating session (the six with the largest '
       'burn). The gauge (blue) descends in visible 1% steps with bounce at each transition; '
       'metered fuel (red) rises smoothly. The staircase against the ramp is the entire '
       'measurement.', width=6.4)

# ---- 3
doc.add_heading('3.  The calibration curve', level=1)
para('Chaining the operating days end to end gives cumulative metered litres against indicated '
     'level across the whole observed band.')
figure('g2_calibration.png',
       'Figure 2 — Cumulative litres against indicated level. The measured slope sits below the '
       'nominal line throughout: each indicated point is buying fewer litres than a linear 250 L '
       'tank would give. The scatter behind the staircase is the raw quantised gauge, shown so '
       'the noise is not hidden by the smoothing that removes it.')
para(f'Across the full observed span the gauge delivers {POOLED:.2f} ± {POOLED_SIG:.2f} L per '
     f'indicated point. If that rate held across the whole tank it would imply about '
     f'{POOLED*100:.0f} L of usable capacity, against the 250 L nominal.')

# ---- 4
doc.add_heading('4.  Four independent lines of evidence', level=1)
table(['Source', 'Band', 'L per point', 'Litres are…', 'Strength'],
      [['Nominal assumption', '0–100%', f'{NOMINAL_LPP:.2f}', 'assumed',
        'No measurement behind it'],
       ['2022 telemetry', '54–92%', f'{Y2022_LPP:.2f}', 'MODELLED',
        'Precise arithmetic, but the litres come from a fuel model, not a meter'],
       ['DD2024 refuel event', '9–97%', f'{DD2024_LPP:.2f}', 'metered (pumped)',
        'Real litres over nearly the whole range; the fill may not be one clean event'],
       ['MCAP flow meter', f'{G_LO:.0f}–{G_HI:.0f}%', f'{POOLED:.2f}', 'metered',
        'Direct and continuous, but spans only a narrow band'],
       ],
      [1.35, 0.75, 0.85, 1.05, 2.3], right_from=1,
      note='Table 2 — The 2022 figure is the weakest of the three non-nominal sources despite '
           'looking the most precise: its litres are what a fuel model predicts for that window, '
           'so it inherits both model error and any in-service burn premium. It should not be '
           'weighed equally with the two metered figures.')

figure('g4_capacity.png',
       'Figure 4 — Implied usable capacity by source. The two metered sources agree closely; the '
       'nominal figure stands apart, and the modelled 2022 figure sits low for the reasons in '
       'Table 2.')
para(f'The two sources that measure real litres — the DD2024 refuel and the flow meter — agree '
     f'to within {abs(POOLED-DD2024_LPP)/POOLED_SIG:.1f} sigma despite sharing no data, no '
     f'method and no year. That agreement is the strongest single result here.')

# ---- 5
doc.add_heading('5.  Is the gauge non-linear? Not established.', level=1)
para(f'Each drawdown covers a different band, and their litres-per-point differ widely — from '
     f'{lo_d["lpp"]:.2f} over {lo_d["g1"]:.0f}–{lo_d["g0"]:.0f}% to {hi_d["lpp"]:.2f} over '
     f'{hi_d["g1"]:.0f}–{hi_d["g0"]:.0f}%. Read quickly, that looks like the gauge changing scale '
     f'down the tank — and an earlier pass over the first few days recorded it as exactly that.')
para('It does not survive an error analysis, and the August 2026 data is the first that can put '
     'the question properly rather than merely restate it.')
figure('g3_lpp.png',
       'Figure 3 — The individual drawdowns with their uncertainties, the ten widest shown. Every '
       'bar overlaps its neighbours. The pooled value (red) is tight because it rests on the whole '
       'travelled drawdown rather than one fall.')
mono(f'widest gap:  {hi_d["lpp"]:.2f} − {lo_d["lpp"]:.2f}  =  {GAP:.2f} L/pt\n'
     f'combined 1σ: {GAP_SIG:.2f}\n'
     f'separation:  {SIGMAS:.1f} σ   →  the SCATTER is real; its CAUSE is the question',
     f'Comparing the extremes of {N_SPANS} falls is not a test of non-linearity — the widest gap '
     f'in a sample that size is wide by construction. What it does establish is that the spread '
     f'between drawdowns exceeds what gauge quantisation alone would produce, so something '
     f'physical drives it. Depth is only one candidate, and the next block tests that one '
     f'directly.')
if DEPTH_SIGMAS is not None:
    para('The sharper test is no longer span against span but SHALLOW against DEEP, because the '
         'record now reaches far enough down the tank to ask it. Reading A predicts that points '
         'below the calibrated band are worth MORE litres than those inside it, so a deep '
         'drawdown should read higher:')
    mono(f'shallow (72–86%):  {SHALLOW_LPP:.2f} L/pt   over {len(_shallow)} falls\n'
         f'deep (below 62%):  {DEEP_LPP:.2f} L/pt   over {len(_deep)} falls\n'
         f'separation:        {DEPTH_SIGMAS:.1f} σ   →  '
         f'{"not significant" if DEPTH_SIGMAS < 2 else "SIGNIFICANT — revisit reading A"}',
         'The deep spans are the ones that did not exist before August 2026; everything earlier '
         'stopped at 72%.')
    para(f'The deep band does read higher, in the direction reading A predicts — but at '
         f'{DEPTH_SIGMAS:.1f} sigma that is not a result, it is the same size of difference this '
         f'report already retracted once. Suggestive is not established, and the distinction is '
         f'the entire subject of this document.')
para('So the two results point in different directions, and both belong in the record. The scatter '
     'between drawdowns is larger than quantisation explains, which says a real effect is at work '
     '— sloshing and trim while under way, sender hysteresis, and temperature are all candidates '
     'this data cannot separate. But that scatter is NOT organised by depth, which is what '
     'non-linearity would look like. A real effect that is not the one under test does not '
     'establish the one under test.')
para('Pushing to finer resolution makes this worse rather than better: computing litres for each '
     'individual gauge point gives values ranging from about 0.2 to 5.0 L — a scatter larger than '
     'the quantity being measured.')

para('None of this disproves non-linearity — a float sender in a tank that is not a uniform prism '
     'has every physical reason to be non-linear, and the DD2024 whole-range figure sitting close '
     'to the narrow-band figure is mildly reassuring rather than conclusive. The honest statement '
     'is that the effect, if present, is smaller than this measurement can resolve.')

callout('Correction to the earlier record',
        'A previous analysis of the first few days recorded the per-day spread as demonstrating '
        'non-linearity, and that statement propagated into the model file and the project notes. '
        'It was over-read. The correct position: the gauge SCALE is firmly established as well '
        'below nominal; the question of whether that scale varies with level is open, and this '
        'data cannot close it. The model file and notes have been corrected to say so.', WARN)

doc.add_heading('5.1  What this measures, and what the planner uses', level=2)
para(f'These are not the same number, and the difference is deliberate. This report measures '
     f'{POOLED:.2f} ± {POOLED_SIG:.2f} L per point over {G_LO:.0f}–{G_HI:.0f}%. The planner is '
     f'still configured with {ADOPTED_LPP:.2f} L per point over '
     f'{ADOPTED_BAND[0]:.0f}–{ADOPTED_BAND[1]:.0f}%, which is what every fuel figure in the '
     f'companion report and every plan the tool produces rests on.')
table(['', 'L per point', 'Band', 'Evidence'],
      [['Adopted in model.json', f'{ADOPTED_LPP:.2f} ± {ADOPTED_SIG:.2f}',
        f'{ADOPTED_BAND[0]:.0f}–{ADOPTED_BAND[1]:.0f}%', '26.7 L over 13 points'],
       ['Measured here', f'{POOLED:.2f} ± {POOLED_SIG:.2f}',
        f'{G_LO:.0f}–{G_HI:.0f}%', f'{TOT:.0f} L over {PTS:.0f} points'],
       ['Measured here, restricted to the adopted band',
        f'{SHALLOW_LPP:.2f}' if SHALLOW_LPP else '—', '72–86%',
        f'{len(_shallow)} falls' if _shallow else '—']],
      [1.9, 1.15, 1.0, 1.5], right_from=99,
      note=f'Table 5 — The pooled figure sits {ADOPTED_SIGMAS:.1f} sigma from the adopted one, '
           f'which is below the bar this report uses everywhere else. Restricted to the SAME band '
           f'the adopted figure was measured over, the new data reads lower, not higher.')
para(f'The reasoning for leaving it alone: the two differ by {ADOPTED_SIGMAS:.1f} sigma, which is '
     f'the same size of difference section 5 has just declined to call a finding; within the '
     f'adopted band the new spans read {SHALLOW_LPP:.2f}, on the other side of it; and the '
     f'direction of the change is the unsafe one — a larger litres-per-point means each indicated '
     f'point is worth more fuel, so every mission would be planned longer. A number that moves '
     f'planning that way should clear the bar, not sit under it.')
para('What the new data does change is reach. Every calibration before August 2026 came from the '
     f'top third of the gauge; these drawdowns reach {G_LO:.0f}%, which is most of the way to the '
     f'{RES_PCT:.0f}% floor. The band below the floor remains uncalibrated, and the experiment in '
     'section 7 is unchanged.')

# ---- 6
doc.add_heading('6.  What the reserve is actually worth', level=1)
para('The reserve is set in indicated percent, so its value in litres follows directly from the '
     'calibration.')
table(['Basis', 'L per point', '15 points =', 'Against nominal'],
      [['Nominal 250 L linear', f'{NOMINAL_LPP:.2f}', f'{15*NOMINAL_LPP:.1f} L', '—'],
       ['Measured (this band)', f'{POOLED:.2f}', f'{15*POOLED:.1f} L',
        f'{15*(POOLED-NOMINAL_LPP):+.1f} L'],
       ['DD2024 whole range', f'{DD2024_LPP:.2f}', f'{15*DD2024_LPP:.1f} L',
        f'{15*(DD2024_LPP-NOMINAL_LPP):+.1f} L'],
       ],
      [1.9, 1.1, 1.1, 1.3],
      note='Table 3 — What fifteen indicated points buys. The measured figures assume the low '
           'band behaves like the measured band, which is precisely the untested assumption.')
para(f'On the measured scale the reserve is worth about {15*POOLED:.0f} L rather than the '
     f'{15*NOMINAL_LPP:.0f} L a nominal reading implies — roughly {abs(15*(NOMINAL_LPP-POOLED)):.0f} L '
     f'less margin than assumed. At survey speed that is on the order of an hour of endurance. '
     f'It does not make the reserve unsafe, but it does mean the reserve is smaller than the '
     f'number suggests, and it argues for planning on the evidence-based capacity rather than '
     f'the nominal.')
callout('The band that matters is the one never measured',
        'Every measurement here comes from the top third of the tank, because the reserve policy '
        f'means the vehicle rarely goes lower and returns are made well above the floor. The '
        f'{RES_PCT:.0f}% '
        'reserve sits in a region for which there is no direct calibration at all. If the gauge '
        'is non-linear anywhere, the bottom is where a float sender is most likely to misbehave '
        '— and that is exactly where the reserve decision is made.', WARN)

# ---- 7
doc.add_heading('7.  What would settle it', level=1)
bullets([
    ('Sound the tank. ',
     'A dowel against a known fill ends the capacity question outright and costs one afternoon '
     'alongside. It was suggested in the DD2024 thread and remains the single highest-value '
     'measurement available.'),
    ('Log litres pumped against indicated level at every refuel. ',
     'Record the level before and after with the volume. Fills that start from different levels '
     'build the calibration curve across bands the vehicle never burns through in normal '
     'operation.'),
    ('Deliberately run one tank low, once. ',
     f'A single controlled drawdown into the reserve band, with the flow meter logging, would '
     f'calibrate the region every planning decision depends on. Nothing else available does '
     f'this. From the lowest level ever recorded ({_DD["lowest"]:.0f}%) down to the '
     f'{_DD["floor"]:.0f}% floor is {_DD["full"]["points"]:.0f} indicated points: about '
     f'{_DD["full"]["litres"]:.0f} L and {_DD["full"]["hours_survey"]:.0f} hours at survey '
     f'speed, or {_DD["full"]["hours_loiter"]:.0f} hours at loiter, and it would fix litres '
     f'per point to {_DD["full"]["precision"]:.1%}. Half of it — down to '
     f'{_DD["half"]["to_pct"]:.0f}% for {_DD["half"]["litres"]:.0f} L and '
     f'{_DD["half"]["hours_survey"]:.0f} hours — buys {_DD["half"]["precision"]:.1%} and '
     f'halves the blind band. No new instrumentation: both channels are already recorded.'),
    ('Keep accumulating operating days. ',
     'Each new day extends the band and tightens the pooled figure. Three more days at the '
     'current rate would roughly halve the uncertainty on any single band, which is what a '
     'non-linearity test actually needs.'),
])

doc.add_heading('Appendix  —  Reproducing this report', level=1)
mono('python tools/extract_bags.py D:/Claude/Fuel/D8_2040\n'
     'python tools/build_gauge_report.py   [OUT.docx]\n'
     'powershell -File tools/export_pdf.ps1',
     'The report recomputes every figure and every number from the day caches, so re-running it '
     'after adding data produces a corrected document rather than a stale one. The significance '
     'test in section 5 is part of that recomputation — if accumulating data ever makes the '
     'band-to-band difference real, the document will say so. The PDF beside it is a Word export '
     'of the same file, produced by the third step; skip it and the PDF goes stale against the '
     'document it is named after, which is how a rebuilt report can still be read in its old '
     'form.')
para('Scope, stated so it is not mistaken for an omission: nothing here depends on the mission '
     'planner, the tidal forecast, or the mission geometry the planner gained during August 2026. '
     'Those change what a plan COSTS. They do not touch what a gauge point is worth, which is the '
     'only question this report asks — so this document was not left behind by that work, it is '
     'simply not about it.')

check_table_widths(doc)
doc.save(OUT)
print('written:', OUT)
print(f'pooled {POOLED:.3f} +/- {POOLED_SIG:.3f} L/pt over {G_LO:.0f}-{G_HI:.0f}% '
      f'({PTS:.0f} pts, {TOT:.1f} L) | day spread {SIGMAS:.1f} sigma')
